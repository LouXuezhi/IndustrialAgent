import uuid
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any, Tuple

# 数据库依赖
from sqlalchemy.ext.asyncio import AsyncSession

# LangChain & 向量库依赖
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from vectorDB.Chroma import chroma_db,embedding_model

# 文件解析依赖
import fitz  # PyMuPDF
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openpyxl import load_workbook



@dataclass
class IngestionReport:
    document_id: uuid.UUID
    chunk_count: int
    # 扩展：增加状态字段以便调试
    status: str = "success"
    error: str = None


class IndustrialDocLoader:
    """
    [来自 ingestion_test.py]
    多格式文档加载器：负责提取文本、章节信息和页码
    """

    def __init__(self, file_path: Path):
        self.file_path = file_path
        self.ext = self.file_path.suffix.lower()
        self.file_name = self.file_path.name

    def load(self) -> List[Tuple[str, Dict[str, Any]]]:
        """返回：[(章节文本, 章节元数据), ...]"""
        loaders = {
            ".pdf": self._load_pdf_with_chapter,
            ".docx": self._load_docx_with_chapter,
            ".doc": self._load_docx_with_chapter,  # 尝试处理doc
            ".xlsx": self._load_xlsx_with_chapter,
            ".xls": self._load_xlsx_with_chapter,
            ".html": self._load_html_with_chapter,
            ".htm": self._load_html_with_chapter,
            ".txt": self._load_txt_with_chapter,
            ".md": self._load_txt_with_chapter,
        }
        if self.ext not in loaders:
            raise ValueError(f"不支持的格式：{self.ext}")
        return loaders[self.ext]()

    def _load_pdf_with_chapter(self):
        chapters = []
        # 提取PDF大纲（章节）
        pdf_fitz = fitz.open(self.file_path)
        try:
            outline = pdf_fitz.get_toc()  # [(层级, 标题, 页码), ...]
            outline_dict = {entry[2] - 1: entry[1] for entry in outline}  # 页码转0索引
        except Exception:
            outline_dict = {}  # 无大纲处理

        pdf_fitz.close()

        # 按页提取文本，关联章节
        with pdfplumber.open(self.file_path) as pdf:
            current_chapter = "默认章节"
            chapter_text = []
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if not text.strip():
                    continue
                # 切换章节（遇到新大纲标题）
                if page_num in outline_dict:
                    if chapter_text:
                        chapters.append(("\n".join(chapter_text), {
                            "file_name": self.file_name,
                            "chapter": current_chapter,
                            "page": page_num,
                            "file_type": "pdf"
                        }))
                    current_chapter = outline_dict[page_num]
                    chapter_text = [text]
                else:
                    chapter_text.append(text)
            # 保存最后一个章节
            if chapter_text:
                chapters.append(("\n".join(chapter_text), {
                    "file_name": self.file_name,
                    "chapter": current_chapter,
                    "page": len(pdf.pages),
                    "file_type": "pdf"
                }))
        return chapters

    def _load_docx_with_chapter(self):
        chapters = []
        doc = DocxDocument(self.file_path)
        current_chapter = "默认章节"
        chapter_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Heading1/2/3 作为章节标题
            if para.style.name.startswith("Heading"):
                if chapter_text:
                    chapters.append(("\n".join(chapter_text), {
                        "file_name": self.file_name,
                        "chapter": current_chapter,
                        "file_type": "docx"
                    }))
                current_chapter = text
                chapter_text = []
            else:
                chapter_text.append(text)
        if chapter_text:
            chapters.append(("\n".join(chapter_text), {
                "file_name": self.file_name,
                "chapter": current_chapter,
                "file_type": "docx"
            }))
        return chapters

    def _load_xlsx_with_chapter(self):
        chapters = []
        wb = load_workbook(self.file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_text = []
            for row in ws.iter_rows(values_only=True):
                # 将 None 转换为空字符串，并连接
                row_text = [str(cell) if cell is not None else "" for cell in row]
                if any(row_text):
                    sheet_text.append("\t".join(row_text))
            if sheet_text:
                chapters.append(("\n".join(sheet_text), {
                    "file_name": self.file_name,
                    "chapter": sheet_name,
                    "file_type": "xlsx"
                }))
        wb.close()
        return chapters

    def _load_html_with_chapter(self):
        chapters = []
        with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        current_chapter = "默认章节"
        chapter_text = []
        for tag in soup.find_all(["h1", "h2", "h3", "p", "div"]):
            text = tag.get_text(strip=True)
            if not text:
                continue
            if tag.name in ["h1", "h2", "h3"]:
                if chapter_text:
                    chapters.append(("\n".join(chapter_text), {
                        "file_name": self.file_name,
                        "chapter": current_chapter,
                        "file_type": "html"
                    }))
                current_chapter = text
                chapter_text = []
            else:
                chapter_text.append(text)
        if chapter_text:
            chapters.append(("\n".join(chapter_text), {
                "file_name": self.file_name,
                "chapter": current_chapter,
                "file_type": "html"
            }))
        return chapters

    def _load_txt_with_chapter(self):
        chapters = []
        with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.read().split("\n")
        current_chapter = "默认章节"
        chapter_text = []
        title_pattern = re.compile(r"^(#+)\s*(.+)")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            match = title_pattern.match(line)
            if match:
                if chapter_text:
                    chapters.append(("\n".join(chapter_text), {
                        "file_name": self.file_name,
                        "chapter": current_chapter,
                        "file_type": "txt"
                    }))
                current_chapter = match.group(2)
                chapter_text = []
            else:
                chapter_text.append(line)
        if chapter_text:
            chapters.append(("\n".join(chapter_text), {
                "file_name": self.file_name,
                "chapter": current_chapter,
                "file_type": "txt"
            }))
        return chapters


class DocumentIngestor:
    """
    负责文档的解析、切分和入库
    """
    def __init__(self):
        # 1. 初始化嵌入模型
        # 注意：这里使用了 raw string r"..." 避免路径转义问题
        self.embedding_model = embedding_model
        # 2. 初始化向量数据库 (Chroma)
        self.vector_db = chroma_db
        # 3. 初始化文本切分器 (用于二次切分)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=750,
            chunk_overlap=75,
            separators=["\n\n", "\n", ". ", "。 "]
        )

    def _industrial_chunk_process(self, chapters: List[Tuple[str, Dict]]) -> List[Document]:
        """
        [来自 ingestion_test.py 的 industrial_chunk 逻辑]
        执行二次切分策略
        """
        final_chunks = []
        for chap_text, chap_meta in chapters:
            # 短章节直接保留，长章节二次切分
            if len(chap_text) <= self.text_splitter._chunk_size:
                final_chunks.append(Document(page_content=chap_text, metadata=chap_meta))
            else:
                sub_chunks = self.text_splitter.split_text(chap_text)
                for idx, sub_chunk in enumerate(sub_chunks):
                    # 补充分块索引
                    meta = chap_meta.copy()
                    meta["chunk_idx"] = idx
                    final_chunks.append(Document(page_content=sub_chunk, metadata=meta))
        return final_chunks

    async def ingest_file(self, path: Path, session: AsyncSession) -> IngestionReport:
        """
        执行完整摄入流程
        """
        # 占位符：目前 session 未实际使用，但保留以符合接口规范
        _ = session

        # 生成文档 ID
        document_id = uuid.uuid4()

        try:
            # 1. 使用 IndustrialDocLoader 加载并按章节粗分
            loader = IndustrialDocLoader(path)
            chapters = loader.load()

            if not chapters:
                return IngestionReport(document_id=document_id, chunk_count=0, status="warning", error="Empty content")

            # 2. 执行工业级二次切分
            final_docs = self._industrial_chunk_process(chapters)

            # 3. 增强元数据 (注入 document_id 以便追踪)
            for doc in final_docs:
                doc.metadata["document_id"] = str(document_id)
                doc.metadata["source_path"] = str(path)

            # 4. 存入 Chroma 向量库
            # 注意：Chroma 的 add_documents 通常是同步操作，但在 async 函数中调用也是安全的（虽然会阻塞）
            if final_docs:
                self.vector_db.add_documents(final_docs)

            return IngestionReport(
                document_id=document_id,
                chunk_count=len(final_docs),
                status="success"
            )

        except Exception as e:
            # 捕获解析过程中的错误
            return IngestionReport(
                document_id=document_id,
                chunk_count=0,
                status="failed",
                error=str(e)
            )