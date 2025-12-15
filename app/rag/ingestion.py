import re
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Tuple

import chromadb
from chromadb.utils import embedding_functions
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.config import Settings, get_settings
from app.db.models import Chunk, Document


def _extract_text_with_chapters(path: Path) -> list[Tuple[str, dict[str, Any]]]:
    """
    提取文档文本并按章节组织。
    返回: [(章节文本, 章节元数据), ...]
    """
    suffix = path.suffix.lower()
    file_name = path.name
    
    if suffix == ".pdf":
        return _extract_pdf_with_chapters(path, file_name)
    elif suffix in [".docx", ".doc"]:
        return _extract_docx_with_chapters(path, file_name)
    elif suffix in [".md", ".markdown"]:
        return _extract_markdown_with_chapters(path, file_name)
    elif suffix == ".txt":
        return _extract_txt_with_chapters(path, file_name)
    else:
        # 回退到简单文本提取
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
            if text.strip():
                return [(text, {"file_name": file_name, "chapter": "默认章节", "file_type": suffix[1:]})]
            return []
        except Exception:
            return []


def _extract_pdf_with_chapters(path: Path, file_name: str) -> list[Tuple[str, dict[str, Any]]]:
    """提取 PDF 文本，尝试按大纲分章节，否则按页分章节。"""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(path)
        chapters = []
        
        # 尝试提取大纲
        outline_dict = {}
        try:
            if reader.outline:
                # pypdf 的 outline 是嵌套结构
                def extract_outline(outline, page_num_map=None):
                    """递归提取大纲"""
                    if page_num_map is None:
                        # 构建页码映射（简化版，假设每个页面对象有页码信息）
                        page_num_map = {}
                        for i, page in enumerate(reader.pages):
                            # 尝试从页面对象获取页码（pypdf 可能不支持，这里做简化处理）
                            page_num_map[i] = i + 1
                    
                    result = {}
                    for item in outline:
                        if isinstance(item, list):
                            # 嵌套结构
                            result.update(extract_outline(item, page_num_map))
                        else:
                            # 可能是字典或对象
                            if hasattr(item, 'title') and hasattr(item, 'page'):
                                # 某些版本的 pypdf 可能支持
                                page_idx = item.page if isinstance(item.page, int) else 0
                                result[page_idx] = item.title
                    return result
                
                # 简化处理：如果大纲是列表，尝试提取
                if isinstance(reader.outline, list):
                    # 遍历大纲项，尝试获取页码和标题
                    for i, item in enumerate(reader.outline):
                        if hasattr(item, 'title'):
                            # 假设页码是索引+1（简化处理）
                            outline_dict[i] = getattr(item, 'title', f"章节{i+1}")
        except Exception:
            # 大纲提取失败，使用按页切分
            outline_dict = {}
        
        # 按页提取文本
        current_chapter = "默认章节"
        chapter_text = []
        total_chars = 0
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if not page_text.strip():
                continue
            
            total_chars += len(page_text.strip())
            
            # 检查是否有新章节（基于大纲或按页切分）
            if page_num in outline_dict:
                # 遇到新章节，保存当前章节
                if chapter_text:
                    chapters.append((
                        "\n".join(chapter_text),
                        {
                            "file_name": file_name,
                            "chapter": current_chapter,
                            "page": page_num,
                            "file_type": "pdf"
                        }
                    ))
                current_chapter = outline_dict[page_num]
                chapter_text = [page_text]
            else:
                # 如果没有大纲，每 5 页作为一个章节（可配置）
                if not outline_dict and page_num > 0 and page_num % 5 == 0:
                    if chapter_text:
                        chapters.append((
                            "\n".join(chapter_text),
                            {
                                "file_name": file_name,
                                "chapter": current_chapter,
                                "page": page_num,
                                "file_type": "pdf"
                            }
                        ))
                    current_chapter = f"第{page_num // 5 + 1}部分"
                    chapter_text = [page_text]
                else:
                    chapter_text.append(page_text)
        
        # 保存最后一个章节
        if chapter_text:
            chapters.append((
                "\n".join(chapter_text),
                {
                    "file_name": file_name,
                    "chapter": current_chapter,
                    "page": len(reader.pages),
                    "file_type": "pdf"
                }
            ))
        
        # 检查是否为扫描型 PDF
        num_pages = len(reader.pages)
        if num_pages > 0:
            avg_chars_per_page = total_chars / num_pages
            if avg_chars_per_page < 50:
                raise ValueError(
                    f"PDF appears to be scanned/image-based (only {total_chars} characters extracted from {num_pages} pages, "
                    f"average {avg_chars_per_page:.1f} chars/page). "
                    f"Text-based PDFs are required. Please use OCR to convert scanned PDFs to text first."
                )
        
        if not chapters:
            raise ValueError(
                "No text could be extracted from PDF. This may be a scanned/image-based PDF. "
                "Text-based PDFs are required. Please use OCR to convert scanned PDFs to text first."
            )
        
        return chapters
        
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def _extract_docx_with_chapters(path: Path, file_name: str) -> list[Tuple[str, dict[str, Any]]]:
    """提取 DOCX 文本，按 Heading 样式分章节。"""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(path)
        chapters = []
        current_chapter = "默认章节"
        chapter_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查是否为标题样式（Heading1, Heading2, Heading3 等）
            if para.style.name.startswith("Heading"):
                # 遇到新章节，保存当前章节
                if chapter_text:
                    chapters.append((
                        "\n".join(chapter_text),
                        {
                            "file_name": file_name,
                            "chapter": current_chapter,
                            "file_type": "docx"
                        }
                    ))
                current_chapter = text
                chapter_text = []
            else:
                chapter_text.append(text)
        
        # 保存最后一个章节
        if chapter_text:
            chapters.append((
                "\n".join(chapter_text),
                {
                    "file_name": file_name,
                    "chapter": current_chapter,
                    "file_type": "docx"
                }
            ))
        
        return chapters if chapters else [
            ("\n".join(p.text for p in doc.paragraphs if p.text.strip()), {
                "file_name": file_name,
                "chapter": "默认章节",
                "file_type": "docx"
            })
        ]
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def _extract_markdown_with_chapters(path: Path, file_name: str) -> list[Tuple[str, dict[str, Any]]]:
    """提取 Markdown 文本，按标题分章节。"""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        if not text.strip():
            return []
        
        chapters = []
        current_chapter = "默认章节"
        chapter_text = []
        
        # 匹配 Markdown 标题（# 标题 或 ## 标题）
        title_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        
        lines = text.split("\n")
        for line in lines:
            match = title_pattern.match(line)
            if match:
                # 遇到新章节，保存当前章节
                if chapter_text:
                    chapters.append((
                        "\n".join(chapter_text),
                        {
                            "file_name": file_name,
                            "chapter": current_chapter,
                            "file_type": "md"
                        }
                    ))
                current_chapter = match.group(2).strip()
                chapter_text = []
            else:
                chapter_text.append(line)
        
        # 保存最后一个章节
        if chapter_text:
            chapters.append((
                "\n".join(chapter_text),
                {
                    "file_name": file_name,
                    "chapter": current_chapter,
                    "file_type": "md"
                }
            ))
        
        return chapters if chapters else [
            (text, {
                "file_name": file_name,
                "chapter": "默认章节",
                "file_type": "md"
            })
        ]
    except Exception as e:
        raise ValueError(f"Failed to extract text from Markdown: {str(e)}")


def _extract_txt_with_chapters(path: Path, file_name: str) -> list[Tuple[str, dict[str, Any]]]:
    """提取纯文本，尝试识别标题模式。"""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        if not text.strip():
            return []
        
        chapters = []
        current_chapter = "默认章节"
        chapter_text = []
        
        # 尝试识别标题模式（以 # 开头或全大写行）
        lines = text.split("\n")
        for line in lines:
            stripped = line.strip()
            # 识别标题：以 # 开头，或全大写且长度适中
            is_title = (
                stripped.startswith("#") or
                (stripped.isupper() and 3 <= len(stripped) <= 50 and not stripped.endswith("."))
            )
            
            if is_title and chapter_text:
                # 遇到新章节，保存当前章节
                chapters.append((
                    "\n".join(chapter_text),
                    {
                        "file_name": file_name,
                        "chapter": current_chapter,
                        "file_type": "txt"
                    }
                ))
                current_chapter = stripped.lstrip("#").strip()
                chapter_text = []
            else:
                chapter_text.append(line)
        
        # 保存最后一个章节
        if chapter_text:
            chapters.append((
                "\n".join(chapter_text),
                {
                    "file_name": file_name,
                    "chapter": current_chapter,
                    "file_type": "txt"
                }
            ))
        
        return chapters if chapters else [
            (text, {
                "file_name": file_name,
                "chapter": "默认章节",
                "file_type": "txt"
            })
        ]
    except Exception as e:
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")


def _extract_text_from_file(path: Path) -> str:
    """Extract text from various file formats."""
    suffix = path.suffix.lower()
    
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text_parts = []
            total_chars = 0
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    total_chars += len(page_text.strip())
            
            # 检查是否为文本型PDF
            # 如果提取的文本很少（少于每页平均50个字符），可能是扫描型PDF
            num_pages = len(reader.pages)
            if num_pages > 0:
                avg_chars_per_page = total_chars / num_pages
                if avg_chars_per_page < 50:
                    raise ValueError(
                        f"PDF appears to be scanned/image-based (only {total_chars} characters extracted from {num_pages} pages, "
                        f"average {avg_chars_per_page:.1f} chars/page). "
                        f"Text-based PDFs are required. Please use OCR to convert scanned PDFs to text first."
                    )
            
            if not text_parts:
                raise ValueError(
                    "No text could be extracted from PDF. This may be a scanned/image-based PDF. "
                    "Text-based PDFs are required. Please use OCR to convert scanned PDFs to text first."
                )
            
            return "\n".join(text_parts)
        except ValueError:
            # Re-raise ValueError (our custom errors for scanned PDFs)
            raise
        except Exception as e:
            # Other exceptions (file read errors, etc.)
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    elif suffix in [".docx", ".doc"]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception:
            return ""
    elif suffix in [".md", ".markdown"]:
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        # Fallback: try to read as text
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""


@dataclass
class IngestionReport:
    document_id: uuid.UUID
    chunk_count: int
    vectorized: bool = False
    error: str | None = None


def _resolve_chroma_path(vector_uri: str) -> str:
    """Accept formats like chroma://./chroma_store and return filesystem path."""
    prefix = "chroma://"
    if vector_uri.startswith(prefix):
        return vector_uri[len(prefix) :]
    return vector_uri


def _build_embedding_fn(settings: Settings):
    # Prefer configured embedding provider; fallback to default.
    if settings.llm_provider == "dashscope" and settings.dashscope_api_key:
        # DashScope provides an OpenAI-compatible endpoint; use api_base to direct traffic.
        return embedding_functions.OpenAIEmbeddingFunction(
            api_key=settings.dashscope_embedding_api_key or settings.dashscope_api_key,
            model_name=settings.embedding_model,
            api_base=settings.dashscope_embedding_base_url or settings.dashscope_base_url,
        )
    if settings.llm_provider == "openai" and settings.openai_api_key:
        return embedding_functions.OpenAIEmbeddingFunction(
            api_key=settings.openai_api_key,
            model_name=settings.embedding_model,
        )
    return embedding_functions.DefaultEmbeddingFunction()


class DocumentIngestor:
    """Persist documents/chunks and write embeddings into Chroma vector store."""

    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()
        chroma_path = _resolve_chroma_path(self.settings.vector_db_uri)
        self._client = chromadb.PersistentClient(path=chroma_path)
        self._embedding_fn = _build_embedding_fn(self.settings)

    def _get_collection(self, library_id: uuid.UUID | None):
        name = f"library_{library_id}" if library_id else "library_default"
        return self._client.get_or_create_collection(
            name=name,
            embedding_function=self._embedding_fn,
            metadata={"library_id": str(library_id) if library_id else None},
        )

    def _smart_chunk_with_chapters(
        self, path: Path, document_id: uuid.UUID, chunk_size: int
    ) -> list[Chunk]:
        """
        智能切分策略：短章节保留，长章节二次切分。
        基于章节信息进行切分，保留文档结构。
        """
        chunks: list[Chunk] = []
        
        try:
            # 尝试使用章节感知的提取
            chapters = _extract_text_with_chapters(path)
        except Exception:
            # 回退到简单文本提取
            try:
                text = _extract_text_from_file(path)
                if not text.strip():
                    return []
                chapters = [(text, {"file_name": path.name, "chapter": "默认章节", "file_type": path.suffix[1:] if path.suffix else "unknown"})]
            except Exception:
                return []
        
        # 对每个章节进行智能切分
        for chapter_text, chapter_meta in chapters:
            if not chapter_text.strip():
                continue
            
            # 短章节直接保留
            if len(chapter_text) <= chunk_size:
                chunk = Chunk(
                    id=uuid.uuid4(),
                    document_id=document_id,
                    content=chapter_text,
                    meta={
                        "offset": 0,
                        "length": len(chapter_text),
                        "chapter": chapter_meta.get("chapter", "默认章节"),
                        "file_name": chapter_meta.get("file_name", path.name),
                        "file_type": chapter_meta.get("file_type", "unknown"),
                        **({k: v for k, v in chapter_meta.items() if k not in ["chapter", "file_name", "file_type"]}),
                    },
                )
                chunks.append(chunk)
            else:
                # 长章节二次切分，保留章节信息
                # 使用简单的重叠切分策略
                chunk_overlap = min(75, chunk_size // 10)  # 10% 重叠，最多75字符
                
                start = 0
                chunk_idx = 0
                while start < len(chapter_text):
                    end = start + chunk_size
                    part = chapter_text[start:end]
                    
                    # 尝试在句号、换行符处切分，避免截断句子
                    if end < len(chapter_text):
                        # 向后查找合适的切分点
                        for sep in ["\n\n", "\n", "。", ". ", "！", "! "]:
                            last_sep = part.rfind(sep)
                            if last_sep > chunk_size * 0.7:  # 至少保留70%的内容
                                part = part[:last_sep + len(sep)]
                                end = start + len(part)
                                break
                    
                    if part.strip():
                        chunk = Chunk(
                            id=uuid.uuid4(),
                            document_id=document_id,
                            content=part,
                            meta={
                                "offset": start,
                                "length": len(part),
                                "chunk_idx": chunk_idx,
                                "chapter": chapter_meta.get("chapter", "默认章节"),
                                "file_name": chapter_meta.get("file_name", path.name),
                                "file_type": chapter_meta.get("file_type", "unknown"),
                                **({k: v for k, v in chapter_meta.items() if k not in ["chapter", "file_name", "file_type"]}),
                            },
                        )
                        chunks.append(chunk)
                        chunk_idx += 1
                    
                    # 移动到下一个切分点（考虑重叠）
                    start = end - chunk_overlap if end < len(chapter_text) else end
        
        return chunks

    async def vectorize_document(
        self,
        document: Document,
        session: AsyncSession,
        chunk_size: int = 800,
    ) -> IngestionReport:
        """Chunk an existing document file and write embeddings."""
        path = Path(document.source_path)
        if not path.exists():
            error_msg = f"Document file not found on disk: {path}"
            meta = dict(document.meta or {})
            meta["vectorized"] = False
            document.meta = meta
            await session.commit()
            return IngestionReport(document_id=document.id, chunk_count=0, vectorized=False, error=error_msg)

        # Clear existing chunks to avoid duplication on re-run
        existing_chunks = await session.execute(select(Chunk).where(Chunk.document_id == document.id))
        for ch in existing_chunks.scalars().all():
            await session.delete(ch)
        await session.flush()

        # 使用增强的章节感知切分
        try:
            chunks = self._smart_chunk_with_chapters(path, document.id, chunk_size)
        except ValueError as e:
            # PDF text extraction errors (e.g., scanned PDFs)
            error_msg = str(e)
            meta = dict(document.meta or {})
            meta["vectorized"] = False
            document.meta = meta
            await session.commit()
            return IngestionReport(document_id=document.id, chunk_count=0, vectorized=False, error=error_msg)
        except Exception as e:
            error_msg = f"Failed to process document: {str(e)}"
            meta = dict(document.meta or {})
            meta["vectorized"] = False
            document.meta = meta
            await session.commit()
            return IngestionReport(document_id=document.id, chunk_count=0, vectorized=False, error=error_msg)
        
        if not chunks:
            error_msg = f"Could not extract any content from file: {path}"
            meta = dict(document.meta or {})
            meta["vectorized"] = False
            document.meta = meta
            await session.commit()
            return IngestionReport(document_id=document.id, chunk_count=0, vectorized=False, error=error_msg)

        session.add_all(chunks)
        await session.flush()  # Flush to ensure chunk IDs are generated
        # write embeddings (best-effort but report failures)
        vectorized = False
        error: str | None = None
        try:
            collection = self._get_collection(document.library_id)
            ids = [str(chunk.id) for chunk in chunks]
            # Ensure no None IDs
            if None in ids:
                raise ValueError("Some chunk IDs are None after flush")
            documents = [chunk.content for chunk in chunks]
            metadatas: list[dict[str, Any]] = []
            for chunk in chunks:
                meta = {
                    "document_id": str(document.id),
                    "offset": chunk.meta.get("offset"),
                    "length": chunk.meta.get("length"),
                }
                # 添加章节信息到元数据
                if "chapter" in chunk.meta:
                    meta["chapter"] = chunk.meta["chapter"]
                if "chunk_idx" in chunk.meta:
                    meta["chunk_idx"] = chunk.meta["chunk_idx"]
                if "page" in chunk.meta:
                    meta["page"] = chunk.meta["page"]
                if "file_name" in chunk.meta:
                    meta["file_name"] = chunk.meta["file_name"]
                if "file_type" in chunk.meta:
                    meta["file_type"] = chunk.meta["file_type"]
                if document.library_id:
                    meta["library_id"] = str(document.library_id)
                metadatas.append(meta)
            collection.add(ids=ids, documents=documents, metadatas=metadatas)
            vectorized = True
        except Exception as exc:
            vectorized = False
            error = str(exc)

        # mark vectorized flag in meta
        meta = dict(document.meta or {})
        meta["chunk_size"] = chunk_size
        meta["vectorized"] = vectorized
        document.meta = meta
        await session.commit()
        await session.refresh(document)

        return IngestionReport(document_id=document.id, chunk_count=len(chunks), vectorized=vectorized, error=error)

